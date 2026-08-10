"""
Завантаження документів бази знань: .md, .txt, .pdf, .docx.

Принцип: loader віддає **чистий текст із збереженими абзацами і заголовками**.
Уся розумна робота (розбиття, breadcrumbs) — далі, у chunking.py. Тут лише
дістати текст і не втратити структуру.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from src.schema import Document

logger = logging.getLogger(__name__)

SUPPORTED = {".md", ".markdown", ".txt", ".pdf", ".docx"}

_MULTI_BLANK = re.compile(r"\n{3,}")


class UnsupportedFormat(Exception):
    pass


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # три і більше порожніх рядки нічого не додають, але з'їдають бюджет чанка
    return _MULTI_BLANK.sub("\n\n", text).strip()


def _title_from(text: str, fallback: str) -> str:
    """Заголовок = перший H1 у markdown або перший непорожній рядок."""
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            return line[2:].strip()
        return line[:120]
    return fallback


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - залежність у requirements
        raise UnsupportedFormat("Для .pdf потрібен пакет pypdf") from exc

    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        content = (page.extract_text() or "").strip()
        if not content:
            # Скан без текстового шару — мовчазно проковтнути його гірше, ніж
            # сказати вголос: користувач думатиме, що документ проіндексовано.
            logger.warning("%s: сторінка %d без текстового шару (скан?)", path.name, i)
            continue
        # Номер сторінки лишаємо як заголовок секції — цитата без сторінки в PDF
        # марна, її неможливо перевірити.
        pages.append(f"## Сторінка {i}\n\n{content}")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedFormat("Для .docx потрібен пакет python-docx") from exc

    doc = docx.Document(str(path))
    out: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        # Заголовки Word переводимо в markdown — далі структурний чанкер
        # працює з ними так само, як з .md.
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "2"
            out.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            out.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return "\n\n".join(out)


def load_file(path: Path) -> Document:
    ext = path.suffix.lower()
    if ext not in SUPPORTED:
        raise UnsupportedFormat(f"{path.name}: формат {ext} не підтримується")

    if ext == ".pdf":
        raw = _read_pdf(path)
    elif ext == ".docx":
        raw = _read_docx(path)
    else:
        raw = path.read_text(encoding="utf-8")

    text = _clean(raw)
    return Document(
        doc_id=path.stem,
        title=_title_from(text, path.stem),
        source_path=str(path).replace("\\", "/"),
        ext=ext,
        text=text,
    )


def load_corpus(root: Path) -> list[Document]:
    """
    Рекурсивно збирає підтримувані файли. Файли, які не вдалось прочитати, не
    валять інжест — пропускаються з попередженням, решта індексується.
    """
    if root.is_file():
        return [load_file(root)]

    docs: list[Document] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED:
            continue
        try:
            doc = load_file(path)
        except Exception as exc:  # noqa: BLE001 — один битий файл не має валити інжест
            logger.warning("Пропущено %s: %s", path.name, exc)
            continue
        if not doc.text:
            logger.warning("Пропущено %s: порожній текст", path.name)
            continue
        docs.append(doc)
    return docs
